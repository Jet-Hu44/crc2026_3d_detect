"""
生成资格认证模板 .docx 文件
包含: 封面信息页 + 第二部分(过往参赛证明) + 第三部分(贡献证明) + 第四部分(赛项贡献) + 授权声明
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def add_heading_centered(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(text, bold=False, size=12, align=None, font_name='宋体'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if align is not None:
        p.alignment = align
    return p


def set_cell_font(cell, text, bold=False, size=10.5):
    """设置表格单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ============================================================
# 封面
# ============================================================

# 标题
add_para('2026 中国机器人大赛暨 RoboCup 机器人世界杯中国赛', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体')
add_para('', size=6)
add_para('机器人先进视觉赛项参赛队资格认证', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体')
add_para('', size=12)

# 封面信息表
cover_data = [
    ('参赛学校', '东北电力大学'),
    ('队伍名称', 'HS（幻视）'),
    ('队伍编号', 'Y2507T1892934'),
    ('参赛队员', '（请填写队员1~5姓名）'),
    ('指导教师\n（姓名/联系方式）', '（请填写教师姓名 / 手机号 / 邮箱）'),
]

table = doc.add_table(rows=len(cover_data), cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(cover_data):
    set_cell_font(table.rows[i].cells[0], label, bold=True, size=11)
    set_cell_font(table.rows[i].cells[1], value, bold=False, size=11)
    table.rows[i].cells[0].width = Cm(4)
    table.rows[i].cells[1].width = Cm(10)

add_para('', size=12)
add_para('参赛项目', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('中国机器人大赛暨 RoboCup 中国赛', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('机器人先进视觉赛项技术委员会', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('2026 年 7 月', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# 第二部分：过往参赛证明
# ============================================================

add_heading_centered('第二部分：过往参赛证明', level=1)
add_para('')
add_para('近 3 年（2023、2024、2025）参加中国机器人大赛机器人先进视觉项目获奖情况说明，同时需提供相应证明材料图片。', size=12)
add_para('')

# 过往参赛表
t2 = doc.add_table(rows=4, cols=5, style='Table Grid')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['年份', '比赛名称', '获奖等级', '证明材料', '备注']
for j, h in enumerate(headers2):
    set_cell_font(t2.rows[0].cells[j], h, bold=True, size=10.5)

# 空行供填写
for i in range(1, 4):
    for j in range(5):
        set_cell_font(t2.rows[i].cells[j], '（请填写）' if i == 1 else '', size=10.5)

add_para('')
add_para('☐ 如为首次参赛，请勾选此处，无需填写上表。请在下方附一份说明：', bold=False, size=11)
add_para('')
add_para('（首次参赛说明请在此处填写...）', size=11)
add_para('')
add_para('')

doc.add_page_break()

# ============================================================
# 第三部分：贡献证明材料
# ============================================================

add_heading_centered('第三部分：贡献证明材料', level=1)
add_para('')
add_para('近 3 年团队或团队成员公开发表的与此机器人涉及技术相关的论文、申请的专利与软件著作权等情况说明，同时需提供相应证明材料图片。', size=12)
add_para('')

# 贡献材料表
t3 = doc.add_table(rows=4, cols=5, style='Table Grid')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['类型（论文/专利/软著）', '名称', '编号/期刊', '日期', '证明材料']
for j, h in enumerate(headers3):
    set_cell_font(t3.rows[0].cells[j], h, bold=True, size=10.5)

for i in range(1, 4):
    for j in range(5):
        set_cell_font(t3.rows[i].cells[j], '（请填写）' if i == 1 else '', size=10.5)

add_para('')
add_para('☐ 如无相关论文、专利或软件著作权，请勾选此处，无需填写上表。请在下方附一份说明：', size=11)
add_para('')
add_para('（说明文档请在此处填写...）', size=11)

doc.add_page_break()

# ============================================================
# 第四部分：对赛项比赛的贡献
# ============================================================

add_heading_centered('第四部分：对赛项比赛的贡献', level=1)
add_para('')
add_para('对 2025/2026 先进视觉赛项比赛的贡献说明（技术委员会将根据贡献的实际价值决定是否给予加分，最高 20 分）：', size=12)
add_para('')

contributions = [
    ('1. 开源代码贡献',
     '本项目基于 GitHub 开源参考项目进行了全面的功能扩展和规则适配。团队计划在比赛结束后，将自主开发的 OCR 模块、两轮自适应框架、Table 编号等改进以开源方式回馈社区，降低后续参赛队伍的开发门槛。'),
    ('2. 技术方案文档化',
     '团队在开发过程中积累了详细的技术文档（项目分析报告、NPU 部署指南、数据采集指南等），系统性地记录了在 OrangePi AI Pro + ORBBEC Astra Pro Plus 平台上搭建 3D 识别系统的完整流程。'),
    ('3. NPU 部署方案探索',
     '团队在研究 Ascend NPU 推理加速方面进行了积极探索，编写了《NPU 部署指南》，记录了 PyTorch → ONNX → ATC → Ascend OM 的完整转换流程及常见问题解决方案。'),
    ('4. 比赛规则反馈',
     '团队在深入研读比赛规则的过程中，整理了规则解读要点和功能对照清单，将通过 QQ 群（1027375571）向技术委员会反馈规则建议。'),
    ('5. QQ 社区建设',
     '团队积极参与先进视觉赛技术交流 QQ 群（1027375571）的讨论，与其他参赛队伍交流技术问题和解决方案，共建赛项技术社区。'),
]

for title, detail in contributions:
    p = doc.add_paragraph()
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.name = '宋体'
    run_title.font.size = Pt(12)
    run_title.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p2 = doc.add_paragraph()
    run_detail = p2.add_run(detail)
    run_detail.font.name = '宋体'
    run_detail.font.size = Pt(12)
    run_detail.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    add_para('')

add_para('补充贡献（请在此处填写其他贡献内容）：', bold=True, size=12)
add_para('')
add_para('（请填写...）', size=11)

doc.add_page_break()

# ============================================================
# 授权声明 + 签名页
# ============================================================

add_heading_centered('关于技术报告使用授权的说明', level=1)
add_para('')

auth_text = (
    '本人完全了解 2026 中国机器人大赛暨 RoboCup 中国赛关于保留、使用技术报告和研究论文的规定，'
    '即：参赛作品著作权归参赛者本人所有，比赛组委会可以在相关主页上收录并公开参赛作品的设计方案、'
    '技术报告以及参赛模型的视频、图像资料，并将相关内容编纂收录在组委会出版论文集中。'
)
add_para(auth_text, size=12)
add_para('')
add_para('')
add_para('')

# 签名区
sig_data = [
    ('参赛队员签名', '____________________'),
    ('', '____________________'),
    ('', '____________________'),
    ('', '____________________'),
    ('', '____________________'),
    ('带队教师签名', '____________________'),
    ('日　　期', '____________________'),
]

t_sig = doc.add_table(rows=len(sig_data), cols=2)
t_sig.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(sig_data):
    set_cell_font(t_sig.rows[i].cells[0], label, bold=bool(label), size=12)
    set_cell_font(t_sig.rows[i].cells[1], value, bold=False, size=12)
    t_sig.rows[i].cells[0].width = Cm(5)
    t_sig.rows[i].cells[1].width = Cm(9)

add_para('')
add_para('')

# 提交提醒
add_para('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━', size=10)
add_para('⚠ 提交提醒', bold=True, size=11)
add_para('• 压缩包统一命名为：NEEPU_3D识别_资格认证材料.zip', size=10)
add_para('• 压缩包内包含：① 本技术认证文档（PDF格式） ② 视频文件夹（功能展示视频）', size=10)
add_para('• 压缩包大小严格控制在 40MB 以内，超出扣 20 分', size=10)
add_para('• 提交邮箱：403993844@qq.com', size=10)
add_para('• 必须在报名的同时提交，不提交则不具备比赛资格', size=10)
add_para('• 第一部分（队伍介绍+视频+软硬件介绍）为必交材料，不提交则无比赛资格', size=10, bold=True)
add_para('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━', size=10)

# ── 保存 ──
output_path = os.path.join(os.path.dirname(__file__), '资格认证_填写模板.docx')
doc.save(output_path)
print(f'[OK] Template saved to: {output_path}')
print(f'      File size: {os.path.getsize(output_path) / 1024:.1f} KB')
