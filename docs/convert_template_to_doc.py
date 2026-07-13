"""
将 资质认证模板.pdf 原样导出为 .doc 文件
基于 PDF 提取的 8 页内容，忠实地还原排版
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 全局默认样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)  # 小四
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 页边距 ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def add_run_to_paragraph(para, text, font_name='宋体', size=12, bold=False):
    """添加格式化文本到段落"""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def new_para(text='', font_name='宋体', size=12, bold=False, alignment=None, spacing=1.5):
    """创建新段落"""
    p = doc.add_paragraph()
    if text:
        add_run_to_paragraph(p, text, font_name, size, bold)
    pf = p.paragraph_format
    pf.line_spacing = spacing
    if alignment is not None:
        pf.alignment = alignment
    return p


def new_empty_lines(n=1):
    """添加空行"""
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        add_run_to_paragraph(p, '', '宋体', 6)


def set_cell(cell, text, font_name='宋体', size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """设置表格单元格内容"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 垂直居中
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ================================================================
# 第 1 页：封面
# ================================================================

# 顶部标注
p = new_para('附件：参赛队伍资格认证模板', '宋体', 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.5)
new_empty_lines(1)

# 主标题
p = new_para('2026中国机器人大赛暨 RoboCup 机器人世界杯中国赛', '黑体', 16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
new_para('', size=6)
p = new_para('机器人先进视觉赛项参赛队资格认证', '黑体', 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)

new_empty_lines(3)

# 封面信息表
cover_items = [
    ('参 赛 学 校', ''),
    ('队 伍 名 称', ''),
    ('参 赛 队 员', ''),
    ('指 导 教 师\n（姓名/联系方式）', ''),
]

table_cover = doc.add_table(rows=len(cover_items), cols=2)
table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(cover_items):
    set_cell(table_cover.rows[i].cells[0], label, '宋体', 12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(table_cover.rows[i].cells[1], value, '宋体', 12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    # 设置列宽
    table_cover.rows[i].cells[0].width = Cm(4.5)
    table_cover.rows[i].cells[1].width = Cm(10)
    # 设置行高
    tr = table_cover.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="500" w:hRule="atLeast"/>')
    trPr.append(trHeight)

new_empty_lines(2)

# 参赛项目
p = new_para('参 赛 项 目', '黑体', 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
new_empty_lines(1)
p = new_para('中国机器人大赛暨 RoboCup中国赛', '宋体', 12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
p = new_para('机器人先进视觉赛项技术委员会', '宋体', 12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
p = new_para('2026年3月', '宋体', 12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)

doc.add_page_break()

# ================================================================
# 第 2 页：资格认证材料提交说明（上）
# ================================================================

p = new_para('资格认证材料提交说明', '黑体', 16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)
new_para('')

# 一、技术认证文档要求
p = new_para('一、技术认证文档要求', '黑体', 14, bold=True, spacing=1.5)

# 特别注意
p = new_para('')
add_run_to_paragraph(p, '特别注意：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '每支报名的参赛队伍必须在报名的同时提交资格认证材料到指定邮箱（403993844@qq.com），不提交资格认证材料的队伍不具备比赛资格；资格认证材料内容包括三个部分（', '宋体', 12)
add_run_to_paragraph(p, '着重声明：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '资格认证材料中必须包含第一部分，如果提交的材料没有第一部分，不能获得比赛资格）：', '宋体', 12)

new_para('')

# 第一部分
p = new_para('第一部分：必须提交材料', '黑体', 12, bold=True, spacing=1.5)
new_para('')

p = new_para('')
add_run_to_paragraph(p, '①队伍介绍', '宋体', 12, bold=True)
add_run_to_paragraph(p, '，主要包括成员介绍，以前的参赛介绍等，正文字体为宋体小四，1.5倍行距，应尽量保证排版美观且不少于4页。', '宋体', 12)

new_para('')

p = new_para('')
add_run_to_paragraph(p, '②机器人功能展示视频', '宋体', 12, bold=True)
add_run_to_paragraph(p, '，时长应在2分钟到3分钟之间，主要内容为：', '宋体', 12)

items_video = [
    '目标台不同随机背景下的识别',
    '目标台与相机不同距离下的识别',
    '实物与贴纸的辨识',
    '运动中物品的识别',
]
for item in items_video:
    p = new_para(f'    ➢ {item}', '宋体', 12, spacing=1.5)

new_para('')

p = new_para('')
add_run_to_paragraph(p, '③参赛软硬件系统介绍相关材料', '宋体', 12, bold=True)
add_run_to_paragraph(p, '，特别强调，技术委员会关注各参赛队队员的自我创新，不能抄袭，不能与他队雷同，否则有可能被取消比赛资格。主要内容为硬件配置说明；视觉软件界面及功能说明；参赛视觉软件的处理流程、主要算法、测试结果、相关软件技术等，（正文字体为宋体小四，1.5倍行距）应尽量保证排版美观且不少于4页。', '宋体', 12)

new_para('')

# 第二部分
p = new_para('第二部分：过往参赛证明', '黑体', 12, bold=True, spacing=1.5)
new_para('')
p = new_para('近3年（即2023，2024，2025年）参加中国自动化学会组织的中国机器人大赛机器人先进视觉项目的获奖情况说明，同时需提供相应证明材料图片。', '宋体', 12, spacing=1.5)

new_para('')

# 第三部分
p = new_para('第三部分：贡献证明材料', '黑体', 12, bold=True, spacing=1.5)
new_para('')
p = new_para('近3年来团队或团队成员公开发表的与此机器人涉及技术相关的论文、申请的专利与软件著作权等情况说明，同时需提供相应证明材料图片。', '宋体', 12, spacing=1.5)

new_para('')

# 二、技术认证文档评分
p = new_para('二、技术认证文档评分', '黑体', 14, bold=True, spacing=1.5)
new_para('')
p = new_para('技术认证文档评分由技术委员会评定。', '宋体', 12, spacing=1.5)

doc.add_page_break()

# ================================================================
# 第 3 页：评分细则 + 注释
# ================================================================

p = new_para('资格认证材料中必须包含第一部分，如果无法提供其他两部分材料，需提交一份说明文档，对情况予以说明；资格认证材料由先进视觉赛机器人技术委员会进行评分并排序；在比赛成绩出现相同情况下，由资格认证评分来决定队伍排名，资格认证排名靠前的最终比赛排名靠前。', '宋体', 12, spacing=1.5)

new_para('')
p = new_para('资格认证材料评分依据如下：', '宋体', 12, bold=True, spacing=1.5)
new_para('')

p = new_para('')
add_run_to_paragraph(p, '（1）对于必须提交材料：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '此项材料不计分，如果不提交此项材料，直接取消比赛资格；如果提交的材料不合要求，从资格认证总分中扣除相应分数，队伍介绍（扣10分），机器人功能展示视频（扣10分），参赛软硬件系统介绍相关材料（扣10分）。', '宋体', 12)

new_para('')

p = new_para('')
add_run_to_paragraph(p, '（2）对于过往参赛证明材料：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '一项一等奖20分，一项二等奖15分，一项三等奖10分。', '宋体', 12)

new_para('')

p = new_para('')
add_run_to_paragraph(p, '（3）对于贡献证明材料：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '与机器人视觉抓取、物体识别等相关的1篇论文、1项发明专利授权得10分，1项发明专利申请受理、1项软件著作权、1项实用新型专利授权得3分。', '宋体', 12)

new_para('')

# 注释
p = new_para('')
add_run_to_paragraph(p, '注1：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '材料在提交时压缩包统一命名为：XX单位_3D识别_资格认证材料；压缩包内包括一个技术认证文档，及一个视频文件夹。', '宋体', 12)

new_para('')

p = new_para('')
add_run_to_paragraph(p, '注2：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '每队上传的资格认证材料严格控制在40M以内，若大于40M的扣20分。', '宋体', 12)

new_para('')

p = new_para('')
add_run_to_paragraph(p, '注3：', '宋体', 12, bold=True)
add_run_to_paragraph(p, '参赛队在提交资格认证时指出上一年度比赛的贡献，由技术委员会决定是否给予20加分。', '宋体', 12)

# 填充空白到下一页
new_empty_lines(8)

doc.add_page_break()

# ================================================================
# 第 4 页：授权声明 + 签名
# ================================================================

new_empty_lines(2)

p = new_para('关于技术报告使用授权的说明', '黑体', 16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.5)

new_empty_lines(3)

auth_text = (
    '本人完全了解2026中国机器人大赛暨RoboCup中国赛关于保留、'
    '使用技术报告和研究论文的规定，即：参赛作品著作权归参赛者本'
    '人所有，比赛组委会可以在相关主页上收录并公开参赛作品的设计'
    '方案、技术报告以及参赛模型的视频、图像资料，并将相关内容编'
    '纂收录在组委会出版论文集中。'
)
p = new_para(auth_text, '宋体', 12, spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

new_empty_lines(5)

# 签名行
sig_lines = [
    ('参赛队员签名：', '                              '),
    ('带队教师签名：', '                                          '),
    ('日        期：', '                               '),
]
for label, value in sig_lines:
    p = new_para('')
    add_run_to_paragraph(p, label, '宋体', 12, bold=False)
    add_run_to_paragraph(p, value, '宋体', 12)
    p.paragraph_format.line_spacing = 2.0

new_empty_lines(6)

doc.add_page_break()

# ================================================================
# 第 5 页：一、基本信息（必填）
# ================================================================

p = new_para('一、基本信息（必填）', '黑体', 16, bold=True, spacing=1.5)
new_empty_lines(2)

p = new_para('1、队伍介绍', '黑体', 14, bold=True, spacing=1.5)
new_empty_lines(6)

p = new_para('2、机器人功能展示视频介绍', '黑体', 14, bold=True, spacing=1.5)
new_empty_lines(6)

p = new_para('3、参赛软硬件系统介绍', '黑体', 14, bold=True, spacing=1.5)
new_empty_lines(6)

doc.add_page_break()

# ================================================================
# 第 6 页：二、过往参赛证明
# ================================================================

p = new_para('二、过往参赛证明', '黑体', 16, bold=True, spacing=1.5)
new_empty_lines(3)

# 空白表格区域
p = new_para('（请在此处填写近3年参加中国机器人大赛机器人先进视觉项目的获奖情况，并附证明材料图片）', '宋体', 11, spacing=1.5)
new_empty_lines(18)

doc.add_page_break()

# ================================================================
# 第 7 页：三、贡献证明
# ================================================================

p = new_para('三、贡献证明', '黑体', 16, bold=True, spacing=1.5)
new_empty_lines(3)

p = new_para('（请在此处填写近3年团队或团队成员公开发表的论文、专利与软件著作权情况，并附证明材料图片）', '宋体', 11, spacing=1.5)
new_empty_lines(20)

doc.add_page_break()

# ================================================================
# 第 8 页：四、对赛项比赛的贡献
# ================================================================

p = new_para('四、对2025先进视觉赛项比赛的贡献', '黑体', 16, bold=True, spacing=1.5)
new_empty_lines(3)

p = new_para('（请在此处陈述对2025/2026先进视觉赛项的技术贡献、社区贡献或其他形式的贡献）', '宋体', 11, spacing=1.5)

# ── 保存 ──
output_path = os.path.join(os.path.dirname(__file__), '资质认证模板.doc')
doc.save(output_path)
print(f'[OK] .doc file saved to: {output_path}')
print(f'      File size: {os.path.getsize(output_path) / 1024:.1f} KB')
print(f'      Pages: 8')
